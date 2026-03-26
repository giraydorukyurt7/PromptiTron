"""
Advanced Document Understanding System using Gemini 2.5
Supports multiple file formats with AI-powered content extraction and analysis
"""

import logging
from typing import Dict, Any, List, Optional, Union
from pathlib import Path
import mimetypes
import base64
from datetime import datetime

import PIL.Image
import docx

from core.gemini_client import gemini_client
from core.rag_system import rag_system
from core.web_analyzer import web_analyzer

logger = logging.getLogger(__name__)

class DocumentUnderstandingSystem:
    """Advanced document understanding using Gemini 2.5"""
    
    def __init__(self):
        self.supported_formats = {
            'text': ['.txt', '.md', '.rtf'],
            'document': ['.pdf', '.docx', '.doc'],
            'image': ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'],
            'code': ['.py', '.js', '.html', '.css', '.json', '.xml'],
            'web': ['http://', 'https://']  # Web URLs
        }
        
        self.extraction_prompts = {
            'general': """Bu belgeyi analiz et ve aşağıdaki bilgileri çıkar:
1. Ana konu ve içerik özeti
2. Önemli kavramlar ve terimler
3. Eğitim seviyesi (ilkokul, ortaokul, lise, üniversite)
4. Müfredat uyumu (varsa hangi YKS dersi/konusu)
5. Öğrenme hedefleri
6. Zorluk seviyesi
7. Önemli formüller, tanımlar veya kurallar
8. Pratik örnekler ve uygulamalar""",
            
            'educational': """Bu eğitim materyalini YKS perspektifinden analiz et:
1. Hangi YKS dersine ait (Matematik, Fizik, Kimya, Biyoloji, vb.)
2. Hangi üniteler ve konular kapsanıyor
3. Müfredat kazanımları ile uyum
4. Soru tipleri ve çözüm yöntemleri
5. Kavram yanılgıları ve dikkat edilecek noktalar
6. Ön koşul bilgiler
7. İlgili konularla bağlantılar
8. Öğrenci için önerilen çalışma stratejisi""",
            
            'question_analysis': """Bu soru/test materyalini analiz et:
1. Soru türü (çoktan seçmeli, açık uçlu, vb.)
2. Zorluk seviyesi ve YKS uyumu
3. Ölçülmek istenen kazanımlar
4. Çözüm stratejileri ve adımları
5. Yaygın hatalar ve dikkat edilecek noktalar
6. Benzer soru türleri için ipuçları
7. Zaman yönetimi önerileri
8. İlgili konu başlıkları"""
        }
    
    async def process_document(
        self, 
        file_path: Union[str, Path], 
        analysis_type: str = "general",
        custom_prompt: Optional[str] = None
    ) -> Dict[str, Any]:
        """Process document or web URL with AI understanding"""
        try:
            # Check if input is a URL
            input_str = str(file_path)
            if input_str.startswith(('http://', 'https://')):
                return await self._process_web_url(input_str, analysis_type, custom_prompt)
            
            # Process as file
            file_path = Path(file_path)
            
            if not file_path.exists():
                return {"error": "File not found"}
            
            # Extract content based on file type
            content_data = await self._extract_content(file_path)
            
            if content_data.get("error"):
                return content_data
            
            # Analyze content with Gemini
            analysis_result = await self._analyze_content(
                content_data, analysis_type, custom_prompt
            )
            
            # Add to RAG system if successful
            if analysis_result.get("success"):
                await self._add_to_rag(file_path, content_data, analysis_result)
            
            return {
                "success": True,
                "file_info": {
                    "name": file_path.name,
                    "size": file_path.stat().st_size,
                    "type": content_data.get("type"),
                    "format": content_data.get("format")
                },
                "content": content_data.get("content", "")[:500] + "..." if len(content_data.get("content", "")) > 500 else content_data.get("content", ""),
                "analysis": analysis_result.get("analysis"),
                "structured_data": analysis_result.get("structured_data"),
                "educational_metadata": analysis_result.get("educational_metadata")
            }
            
        except Exception as e:
            logger.error(f"Document processing error: {e}")
            return {"error": str(e)}
    
    async def _process_web_url(
        self, 
        url: str, 
        analysis_type: str = "general",
        custom_prompt: Optional[str] = None
    ) -> Dict[str, Any]:
        """Process web URL using web analyzer"""
        try:
            logger.info(f"Processing web URL: {url}")
            
            # Use web analyzer to process the URL
            web_result = await web_analyzer.analyze_website(
                url=url,
                analysis_type=analysis_type,
                custom_prompt=custom_prompt
            )
            
            if web_result.get("error"):
                return {
                    "error": f"Web analizi başarısız: {web_result['error']}",
                    "suggestion": web_result.get("suggestion", "")
                }
            
            # Transform web result to match document processing format
            return {
                "success": True,
                "source_type": "website",
                "url": url,
                "file_info": {
                    "name": web_result.get("content_info", {}).get("title", "Web Sitesi"),
                    "size": web_result.get("content_info", {}).get("word_count", 0),
                    "type": "web",
                    "format": "html"
                },
                "content": "Web içeriği başarıyla analiz edildi",
                "analysis": web_result.get("educational_analysis", ""),
                "structured_data": web_result.get("structured_data", {}),
                "curriculum_check": web_result.get("curriculum_check", {}),
                "generated_questions": web_result.get("generated_questions", []),
                "study_materials": web_result.get("study_materials", {}),
                "educational_metadata": {
                    "yks_relevant": web_result.get("curriculum_check", {}).get("yks_relevant", False),
                    "subjects": web_result.get("curriculum_check", {}).get("subjects", []),
                    "confidence_score": web_result.get("curriculum_check", {}).get("confidence_score", 0.0)
                }
            }
            
        except Exception as e:
            logger.error(f"Web URL processing error: {e}")
            return {"error": f"Web sitesi işleme hatası: {str(e)}"}
    
    async def _extract_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from different file formats"""
        try:
            file_extension = file_path.suffix.lower()
            mime_type = mimetypes.guess_type(str(file_path))[0]
            
            # Text files
            if file_extension in self.supported_formats['text']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                return {
                    "type": "text",
                    "format": file_extension,
                    "content": content,
                    "mime_type": mime_type
                }
            
            # PDF files
            elif file_extension == '.pdf':
                return await self._extract_pdf_content(file_path)
            
            # Word documents
            elif file_extension in ['.docx', '.doc']:
                return await self._extract_docx_content(file_path)
            
            # Images
            elif file_extension in self.supported_formats['image']:
                return await self._extract_image_content(file_path)
            
            # Code files
            elif file_extension in self.supported_formats['code']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                return {
                    "type": "code",
                    "format": file_extension,
                    "content": content,
                    "mime_type": mime_type
                }
            
            else:
                return {"error": f"Unsupported file format: {file_extension}"}
                
        except Exception as e:
            logger.error(f"Content extraction error: {e}")
            return {"error": str(e)}
    
    async def _extract_pdf_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from PDF files"""
        try:
            import pdfplumber
            
            content = ""
            tables = []
            images = []
            
            with pdfplumber.open(file_path) as pdf:
                # Extract text
                for page in pdf.pages:
                    if page.extract_text():
                        content += page.extract_text() + "\n"
                    
                    # Extract tables
                    for table in page.extract_tables():
                        if table:
                            tables.append(table)
                
                # Get basic info
                info = pdf.metadata or {}
            
            return {
                "type": "document",
                "format": ".pdf",
                "content": content,
                "tables": tables,
                "pages": len(pdf.pages),
                "metadata": info
            }
            
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return {"error": f"PDF processing failed: {str(e)}"}
    
    async def _extract_docx_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from Word documents"""
        try:
            doc = docx.Document(file_path)
            
            content = ""
            tables = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                "type": "document",
                "format": ".docx",
                "content": content,
                "tables": tables,
                "sections": len(doc.sections)
            }
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return {"error": f"DOCX processing failed: {str(e)}"}
    
    async def _extract_image_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from images using Gemini vision"""
        try:
            # Load image
            image = PIL.Image.open(file_path)
            
            # Convert to base64 for API
            import io
            buffer = io.BytesIO()
            image.save(buffer, format=image.format or 'PNG')
            image_data = base64.b64encode(buffer.getvalue()).decode()
            
            # Use Gemini vision to extract text and understand content
            vision_prompt = """Bu görüntüyü analiz et ve aşağıdaki bilgileri çıkar:
1. Görüntüde bulunan metin (varsa)
2. Görsel içerik açıklaması
3. Eğitim amaçlı kullanım potansiyeli
4. Hangi derslerde kullanılabilir
5. Görsel öğelerin açıklaması (diyagram, grafik, şekil vb.)
6. Önemli bilgiler ve kavramlar
"""
            
            # Call Gemini with image
            response = await gemini_client.generate_content(
                prompt=vision_prompt,
                images=[image]
            )
            
            return {
                "type": "image",
                "format": file_path.suffix,
                "content": response.get("text", ""),
                "image_info": {
                    "width": image.width,
                    "height": image.height,
                    "format": image.format,
                    "mode": image.mode
                },
                "analysis": response.get("text", "")
            }
            
        except Exception as e:
            logger.error(f"Image extraction error: {e}")
            return {"error": f"Image processing failed: {str(e)}"}
    
    async def _analyze_content(
        self, 
        content_data: Dict[str, Any], 
        analysis_type: str,
        custom_prompt: Optional[str] = None
    ) -> Dict[str, Any]:
        """Analyze extracted content with AI"""
        try:
            content = content_data.get("content", "")
            
            if not content.strip():
                return {"error": "No content to analyze"}
            
            # Select analysis prompt
            if custom_prompt:
                analysis_prompt = custom_prompt
            else:
                analysis_prompt = self.extraction_prompts.get(analysis_type, self.extraction_prompts["general"])
            
            # Combine prompt with content
            full_prompt = f"{analysis_prompt}\n\nBelge İçeriği:\n{content}"
            
            # Analyze with Gemini
            response = await gemini_client.generate_content(
                prompt=full_prompt,
                system_instruction="Sen uzman bir eğitim analisti ve YKS konularında uzmansın. Belgeleri eğitim perspektifinden analiz ediyorsun."
            )
            
            if not response.get("success", True):
                return {"error": "Analysis failed"}
            
            analysis_text = response.get("text", "")
            
            # Extract structured data
            structured_data = await self._extract_structured_info(analysis_text, content)
            
            # Generate educational metadata
            educational_metadata = await self._generate_educational_metadata(content, structured_data)
            
            return {
                "success": True,
                "analysis": analysis_text,
                "structured_data": structured_data,
                "educational_metadata": educational_metadata
            }
            
        except Exception as e:
            logger.error(f"Content analysis error: {e}")
            return {"error": str(e)}
    
    async def _extract_structured_info(self, analysis_text: str, content: str) -> Dict[str, Any]:
        """Extract structured information from analysis"""
        try:
            # Use Gemini to extract structured data
            structured_prompt = f"""
            Analiz metninden yapılandırılmış bilgi çıkar:
            
            Analiz: {analysis_text}
            
            JSON formatında şu bilgileri çıkar:
            {{
                "subject": "ana ders (Matematik, Fizik, Kimya, Biyoloji, vb.)",
                "topics": ["konu1", "konu2", ...],
                "difficulty_level": "kolay/orta/zor",
                "education_level": "lise/üniversite/vb.",
                "key_concepts": ["kavram1", "kavram2", ...],
                "formulas": ["formül1", "formül2", ...],
                "learning_objectives": ["hedef1", "hedef2", ...],
                "exam_relevance": "YKS/TYT/AYT uygunluğu",
                "estimated_study_time": "tahmini çalışma süresi dakika cinsinden"
            }}
            """
            
            response = await gemini_client.generate_content(
                prompt=structured_prompt,
                system_instruction="JSON formatında yanıt ver. Geçerli JSON formatını koru."
            )
            
            try:
                import json
                structured_data = json.loads(response.get("text", "{}"))
                return structured_data
            except json.JSONDecodeError:
                # Fallback to manual parsing
                return {
                    "subject": "Genel",
                    "topics": [],
                    "difficulty_level": "orta",
                    "education_level": "lise",
                    "key_concepts": [],
                    "formulas": [],
                    "learning_objectives": [],
                    "exam_relevance": "YKS",
                    "estimated_study_time": 30
                }
                
        except Exception as e:
            logger.error(f"Structured info extraction error: {e}")
            return {}
    
    async def _generate_educational_metadata(self, content: str, structured_data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate educational metadata for the document"""
        try:
            return {
                "content_length": len(content),
                "reading_time_minutes": max(1, len(content.split()) // 200),  # ~200 words per minute
                "complexity_score": self._calculate_complexity_score(content),
                "curriculum_alignment": structured_data.get("exam_relevance", ""),
                "prerequisite_topics": [],  # Could be enhanced with AI
                "related_topics": structured_data.get("topics", []),
                "study_recommendations": self._generate_study_recommendations(structured_data),
                "assessment_suggestions": self._generate_assessment_suggestions(structured_data)
            }
            
        except Exception as e:
            logger.error(f"Educational metadata generation error: {e}")
            return {}
    
    def _calculate_complexity_score(self, content: str) -> float:
        """Calculate content complexity score (0-1)"""
        try:
            words = content.split()
            
            # Factors affecting complexity
            avg_word_length = sum(len(word) for word in words) / len(words) if words else 0
            sentence_count = content.count('.') + content.count('!') + content.count('?')
            avg_sentence_length = len(words) / sentence_count if sentence_count > 0 else 0
            
            # Technical terms indicators
            technical_indicators = ['formül', 'denklem', 'teoremi', 'kanunu', 'ilkesi', 'integral', 'türev']
            technical_score = sum(1 for indicator in technical_indicators if indicator in content.lower())
            
            # Normalize scores
            word_score = min(avg_word_length / 10, 1)
            sentence_score = min(avg_sentence_length / 20, 1)
            tech_score = min(technical_score / 5, 1)
            
            # Weighted complexity score
            complexity = (word_score * 0.3 + sentence_score * 0.4 + tech_score * 0.3)
            return round(complexity, 2)
            
        except Exception:
            return 0.5  # Default medium complexity
    
    def _generate_study_recommendations(self, structured_data: Dict[str, Any]) -> List[str]:
        """Generate study recommendations based on content analysis"""
        recommendations = []
        
        difficulty = structured_data.get("difficulty_level", "orta")
        subject = structured_data.get("subject", "")
        
        if difficulty == "zor":
            recommendations.extend([
                "Bu konuyu küçük parçalara bölerek çalış",
                "Ön koşul konuları tekrar et",
                "Bol örnek çöz ve pratik yap"
            ])
        elif difficulty == "kolay":
            recommendations.extend([
                "Hızlı geçiş yapabilirsin",
                "Ana kavramları pekiştir",
                "İleri seviye sorulara odaklan"
            ])
        else:
            recommendations.extend([
                "Düzenli tekrar yap",
                "Kavramları anlamaya odaklan",
                "Benzer örnekleri incele"
            ])
        
        if "matematik" in subject.lower():
            recommendations.append("Formülleri ezberlemek yerine anlamaya odaklan")
        elif "fizik" in subject.lower():
            recommendations.append("Günlük yaşam örnekleri ile ilişkilendir")
        
        return recommendations
    
    def _generate_assessment_suggestions(self, structured_data: Dict[str, Any]) -> List[str]:
        """Generate assessment suggestions"""
        suggestions = []
        
        topics = structured_data.get("topics", [])
        subject = structured_data.get("subject", "")
        
        if topics:
            suggestions.append(f"Bu konulardan soru çöz: {', '.join(topics[:3])}")
        
        if structured_data.get("exam_relevance"):
            suggestions.append(f"{structured_data['exam_relevance']} formatında sorular çöz")
        
        suggestions.extend([
            "Konu testi çöz",
            "Zamanlı deneme yap",
            "Hata analizi yap"
        ])
        
        return suggestions
    
    async def _add_to_rag(self, file_path: Path, content_data: Dict[str, Any], analysis_result: Dict[str, Any]):
        """Add processed document to RAG system"""
        try:
            # Convert list values to strings for ChromaDB compatibility
            topics_list = analysis_result.get("structured_data", {}).get("topics", [])
            topics_str = ", ".join(topics_list) if isinstance(topics_list, list) else str(topics_list)
            
            document = {
                "content": content_data.get("content", ""),
                "metadata": {
                    "filename": file_path.name,
                    "file_type": content_data.get("type"),
                    "subject": str(analysis_result.get("structured_data", {}).get("subject", "")),
                    "topics": topics_str,  # Convert list to string
                    "difficulty": str(analysis_result.get("structured_data", {}).get("difficulty_level", "")),
                    "education_level": str(analysis_result.get("structured_data", {}).get("education_level", "")),
                    "upload_date": datetime.now().isoformat(),
                    "analysis": str(analysis_result.get("analysis", ""))[:1000]  # Limit length
                }
            }
            
            await rag_system.add_documents([document], collection_name="documents")
            logger.info(f"Added document {file_path.name} to RAG system")
            
        except Exception as e:
            logger.error(f"Error adding document to RAG: {e}")

# Global document understanding system
document_understanding = DocumentUnderstandingSystem()